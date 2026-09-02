"""Populates data/documents/ with sample docx/csv/pdf files so the RAG app has
something to index out of the box. Safe to re-run; it overwrites its own files."""

from pathlib import Path

import pandas as pd
from docx import Document

DOCS_DIR = Path(__file__).resolve().parent.parent / "data" / "documents"


def make_policy_docx():
    doc = Document()
    doc.add_heading("Company Policy Handbook", level=1)

    doc.add_heading("Refund Policy", level=2)
    doc.add_paragraph(
        "Customers on the Pro Plan or Enterprise Plan may request a full refund "
        "within 14 days of purchase. Starter Plan purchases and Add-on Packs are "
        "non-refundable. Refund requests must be submitted through support and are "
        "processed within 5 business days."
    )

    doc.add_heading("Support SLA", level=2)
    doc.add_paragraph(
        "Enterprise Plan customers receive priority support with a 4-hour response "
        "time. Pro Plan customers receive a 1-business-day response time. Starter "
        "Plan customers are supported via community forums only."
    )

    doc.add_heading("Data Retention", level=2)
    doc.add_paragraph(
        "Customer account data is retained for 90 days after account cancellation, "
        "after which it is permanently deleted. Backups are retained for an "
        "additional 30 days."
    )

    doc.save(DOCS_DIR / "company_policy.docx")


def make_faq_csv():
    rows = [
        ("What plans do you offer?", "We offer Starter, Pro, and Enterprise plans, plus optional Add-on Packs."),
        ("How do I upgrade my plan?", "Go to Account Settings > Billing and select a new plan; changes apply immediately."),
        ("Do you offer discounts for annual billing?", "Yes, annual billing gives a 15% discount versus monthly billing."),
        ("Which regions do you support?", "We currently support customers in APAC, EMEA, and LATAM regions."),
    ]
    pd.DataFrame(rows, columns=["question", "answer"]).to_csv(DOCS_DIR / "product_faq.csv", index=False)


def make_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        print(
            "reportlab not installed - skipping sample PDF. "
            "`pip install reportlab` to generate one, or drop your own PDF into "
            f"{DOCS_DIR}"
        )
        return

    path = DOCS_DIR / "onboarding_guide.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    lines = [
        "Customer Onboarding Guide",
        "",
        "Step 1: Create an account and verify your email address.",
        "Step 2: Choose a plan - Starter, Pro, or Enterprise.",
        "Step 3: Enterprise customers are assigned a dedicated onboarding specialist.",
        "Step 4: All customers receive a welcome email with setup documentation.",
    ]
    y = 750
    for line in lines:
        c.drawString(72, y, line)
        y -= 20
    c.save()


def main():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    make_policy_docx()
    make_faq_csv()
    make_pdf()
    print(f"Sample documents written to {DOCS_DIR}")


if __name__ == "__main__":
    main()
