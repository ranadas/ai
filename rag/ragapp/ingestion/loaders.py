from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class RawDocument:
    source: str
    text: str


def load_pdf(path: Path) -> RawDocument:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return RawDocument(source=path.name, text="\n".join(pages))


def load_docx(path: Path) -> RawDocument:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return RawDocument(source=path.name, text="\n".join(parts))


def load_spreadsheet(path: Path) -> RawDocument:
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)
    lines = [
        ", ".join(f"{col}={row[col]}" for col in df.columns)
        for _, row in df.iterrows()
    ]
    return RawDocument(source=path.name, text="\n".join(lines))


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".csv": load_spreadsheet,
    ".xlsx": load_spreadsheet,
    ".xls": load_spreadsheet,
}


def load_documents(docs_dir: Path) -> list[RawDocument]:
    if not docs_dir.exists():
        return []
    docs = []
    for path in sorted(docs_dir.glob("**/*")):
        if not path.is_file():
            continue
        loader = LOADERS.get(path.suffix.lower())
        if loader is None:
            continue
        try:
            docs.append(loader(path))
        except Exception as exc:
            print(f"Skipping {path.name}: {exc}")
    return docs
